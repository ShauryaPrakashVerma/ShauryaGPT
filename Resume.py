from groq import Groq
from dotenv import load_dotenv
from pydantic import BaseModel
import os
import json
from pypdf import PdfReader
from docx import Document


load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("Invalid API key")

client = Groq(api_key=my_api_key)
model = "llama-3.3-70b-versatile"

file_path = "C:\\Users\\Shaur\\Desktop\\ShauryaGPT\\Personal_Information\\MY PROJECTS.docx"

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

description = read_docx(file_path)

system_prompt = f'''
# ROLE:
You are an expert virtual assistant for Shaurya. You can answer formal conversational questions, like hi, hello etc. Extract information from the resume based on its meaning, not only based on exact section headings.

# TASK:
Your task is to answer all the questions asked related to the projects and details about the person.

# Constraints:
You need to answer about details mentioned in the given description only.
{description }

# Output Format:
The output need to strictly as per the specified format. The output must be easy to interpret by looking at it, the key informations must be preferrably in markdown.

# Fallback
If the question isnt related to the details mentioned in the description, the answer should be just that i can't answer this question. Don't just blindly interpret that questions are not related to the description, also extract it meaning from previous messages and reply accordingly. 

INSTRUCTIONS:
1. DO NOT INVENT DETAILS YOURSELF.
2. BE HONEST WITH THE INFORMATION


'''
# ---------------------------------------------------------------------------------------------------------------------------------------------------------

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


# def 

# ---------------------------------------------------------------------------------------------------------------------------------------------------------







def ask_llm( user_prompt, system_prompt = system_prompt):


    user_prompt = user_prompt

    messages = [
        {
            "role" : "system",
            "content" : system_prompt
        },
        {
            "role" : "user",
            "content" : user_prompt
        }
    ]

    response = client.chat.completions.create(model=model, messages=messages, stream=True)
    # answer = response.choices[0].message.content
    
    answer = ""
    for chunk in response:
        content = chunk.choices[0].delta.content
        if content:
            answer = answer + content
            print(content, end="", flush=True)
        
    messages.append({
        "role" : "assistant",
        "content" : answer
    })
    
    return answer
    
    
def doc_parser():
    pass